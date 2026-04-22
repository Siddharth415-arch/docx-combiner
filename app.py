# -*- coding: utf-8 -*-
"""
Internal Tools — Streamlit Web App
Three tools in one:
  1. DocX Combiner & Formatter  — sorts, merges, and formats chapter headings
  2. Benchmark Converter        — converts docx files to benchmark format
  3. Document Translator        — translates .docx via Gemini AI (upload or Drive link)
"""

import io
import os
import re
import shutil
import tempfile
import threading
import time
import traceback
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import gdown
    GDOWN_AVAILABLE = True
except ImportError:
    GDOWN_AVAILABLE = False


# ══════════════════════════════════════════════════════════════════════
#  SHARED PAGE CONFIG  (must be first Streamlit call)
# ══════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Tools",
    page_icon="🛠️",
    layout="centered",
    menu_items={}  # removes the hamburger "About / Report a bug" menu
)


# ── Hide Streamlit Community Cloud chrome ──
# Specifically targets the "Created by <username>" profile badge and its avatar
# so the app doesn't expose the owner's identity to viewers.
HIDE_STREAMLIT_STYLE = """
<style>
    /* Hamburger menu, footer, header */
    #MainMenu {visibility: hidden !important;}
    footer   {visibility: hidden !important;}
    header[data-testid="stHeader"]  {display: none !important;}
    div[data-testid="stDecoration"] {display: none !important;}
    div[data-testid="stToolbar"]    {display: none !important;}
    .stAppDeployButton              {display: none !important;}
    .stAppHeader                    {display: none !important;}

    /* ── "Created by siddharth415-arch" — profile badge (avatar + username) ── */
    /* Catch every hashed/variable class name used by Streamlit Cloud */
    [class*="profileContainer"]     {display: none !important;}
    [class*="profile_container"]    {display: none !important;}
    [class*="_profileContainer"]    {display: none !important;}
    [class*="_profile_"]            {display: none !important;}
    [class*="stProfileBadge"]       {display: none !important;}
    [class*="ProfileBadge"]         {display: none !important;}
    [data-testid*="stProfile"]      {display: none !important;}
    [data-testid*="profileBadge"]   {display: none !important;}

    /* The <a> tag that wraps the avatar + username and points to the user's profile */
    a[href*="share.streamlit.io/user/"] {display: none !important;}
    a[href*="/user/siddharth"]          {display: none !important;}

    /* Nuclear option — hide any fixed-position element anchored bottom-right
       that contains a link to a Streamlit user profile */
    div:has(> a[href*="share.streamlit.io/user/"]) {display: none !important;}

    /* Also hide Fork / viewer badges in case they're still showing */
    [class*="viewerBadge"]              {display: none !important;}
    [data-testid*="viewerBadge"]        {display: none !important;}
    iframe[title*="viewerBadge"]        {display: none !important;}
</style>
"""
st.markdown(HIDE_STREAMLIT_STYLE, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════
#  SIDEBAR — TOOL SELECTOR
# ══════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("### Select Tool")
    selected_tool = st.radio(
        label="Tool",
        options=["📚 DocX Combiner", "⚙️ Benchmark Converter", "🌐 Document Translator"],
        label_visibility="collapsed",
    )


# ══════════════════════════════════════════════════════════════════════
#  TOOL 1 — DOCX COMBINER & FORMATTER
# ══════════════════════════════════════════════════════════════════════

CHAPTER_PATTERNS = [
    r'^§§§第\d+章.*$',
    r'^第\d+章.*$',
    r'^第[一二三四五六七八九十百千]+章.*$',
    r'^Chapter\s*[-_:]?\s*\d+.*$',
    r'^CHAPTER\s*[-_:]?\s*\d+.*$',
    r'^Ch\.?\s*[-_:]?\s*\d+.*$',
    r'^卷\d+.*$',
    r'^Episode\s*[-_:]?\s*\d+.*$',
    r'^EPISODE\s*[-_:]?\s*\d+.*$',
    r'^Ep\.?\s*[-_:]?\s*\d+.*$',
    r'^EP\.?\s*[-_:]?\s*\d+.*$',     # catches EP1, EP 1, EP-1, EP_1, EP.1
    r'^[Ee][Pp]\s*[-_:]?\s*\d+.*$',  # catches ep1, Ep1, eP1, etc.
    r'^\d+\s*[-–—:.].*$',              # catches "1 - Golden Blood" / "1. Golden Blood"
]


def extract_start_number(filepath: Path) -> int:
    """Pull the episode/chapter number out of the filename stem."""
    name = filepath.stem
    numbers = re.findall(r'\d+', name)
    if not numbers:
        return 0
    nums = [int(n) for n in numbers]
    return min(nums) if len(nums) > 1 else nums[0]


def sort_batch_files(files: list) -> list:
    return sorted(files, key=lambda f: extract_start_number(f))


def unpack_docx(docx_path: Path, unpack_dir: Path):
    unpack_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path, 'r') as z:
        z.extractall(unpack_dir)
    for xml_file in unpack_dir.rglob("*.xml"):
        try:
            tree = etree.parse(str(xml_file))
            with open(xml_file, "wb") as f:
                f.write(etree.tostring(tree, pretty_print=True,
                        encoding="UTF-8", xml_declaration=False))
        except Exception:
            pass


def pack_docx(unpack_dir: Path, output_docx: Path):
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as z:
        for f in unpack_dir.rglob("*"):
            if f.is_file():
                z.write(f, f.relative_to(unpack_dir))


def get_body_paragraphs_raw(doc_xml: Path) -> list:
    try:
        tree = etree.parse(str(doc_xml))
    except Exception:
        return []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = tree.find('.//w:body', ns)
    if body is None:
        return []
    paragraphs = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue
        paragraphs.append(etree.tostring(child, encoding='unicode'))
    return paragraphs


def get_sectPr_raw(doc_xml: Path) -> str:
    try:
        tree = etree.parse(str(doc_xml))
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        sectPr = tree.find('.//w:body/w:sectPr', ns)
        if sectPr is not None:
            return etree.tostring(sectPr, encoding='unicode')
    except Exception:
        pass
    return ''


def get_paragraph_text(para_xml_str: str) -> str:
    """Extract plain text from a raw paragraph XML string."""
    try:
        elem = etree.fromstring(para_xml_str)
    except Exception:
        return ''
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    texts = elem.findall('.//w:t', ns)
    return ''.join(t.text for t in texts if t.text)


def looks_like_heading(text: str) -> bool:
    """True if this paragraph text matches any known chapter/episode pattern."""
    text = text.strip()
    if not text:
        return False
    for pattern in CHAPTER_PATTERNS:
        if re.match(pattern, text, flags=re.IGNORECASE):
            return True
    return False


def extract_first_heading(paragraphs: list, max_check: int = 5) -> tuple:
    """
    Look at the first `max_check` non-empty paragraphs. If any of them looks
    like a chapter/episode heading, return (heading_text, paragraphs_without_it).
    Otherwise return (None, original_paragraphs).
    """
    seen_non_empty = 0
    for i, para_xml in enumerate(paragraphs):
        text = get_paragraph_text(para_xml).strip()
        if not text:
            continue
        seen_non_empty += 1
        if looks_like_heading(text):
            # Remove this paragraph — it'll be rebuilt as an H1
            remaining = paragraphs[:i] + paragraphs[i + 1:]
            return text, remaining
        if seen_non_empty >= max_check:
            break
    return None, paragraphs


def _xml_escape(s: str) -> str:
    return (s.replace('&', '&amp;')
             .replace('<', '&lt;')
             .replace('>', '&gt;'))


def build_heading_paragraph_xml(heading_text: str) -> str:
    """Return raw Heading-1 paragraph XML for the heading text."""
    safe = _xml_escape(heading_text)
    return (
        '<w:p>'
        '<w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        f'<w:r><w:t xml:space="preserve">{safe}</w:t></w:r>'
        '</w:p>'
    )


def derive_heading_from_filename(filepath: Path) -> str:
    """
    Build a chapter/episode heading string from the filename.
    - If the filename contains a number, returns 'Chapter N'.
    - Otherwise uses the filename stem as-is.
    """
    num = extract_start_number(filepath)
    if num > 0:
        return f"Chapter {num}"
    return filepath.stem.strip() or "Chapter"


def merge_docx_files(file_list: list, output_path: Path, progress_callback=None):
    """
    Merge .docx files in the given order. Each file gets EXACTLY ONE Heading 1:
      - If one of the first few paragraphs already looks like a chapter heading,
        we extract that text, remove that paragraph from the body, and rebuild
        it as an H1 at the top (so there's no duplicate heading).
      - Otherwise, we synthesize an H1 from the filename (e.g. 'Chapter 42').
    Returns (from_doc_count, from_filename_count).
    """
    if not file_list:
        raise ValueError("No files to merge.")

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        base_dir = tmp / "base"
        unpack_docx(file_list[0], base_dir)
        base_xml = base_dir / "word" / "document.xml"

        all_paragraphs = []
        from_doc_count      = 0
        from_filename_count = 0

        for i, f in enumerate(file_list):
            if i == 0:
                fxml = base_xml
            else:
                fdir = tmp / f"doc_{i}"
                unpack_docx(f, fdir)
                fxml = fdir / "word" / "document.xml"

            paras = get_body_paragraphs_raw(fxml)

            # Try to find an existing heading in the first few paragraphs
            heading_text, paras = extract_first_heading(paras)

            if heading_text:
                from_doc_count += 1
            else:
                heading_text = derive_heading_from_filename(f)
                from_filename_count += 1

            # Prepend EXACTLY ONE H1 for this chapter/episode
            all_paragraphs.append(build_heading_paragraph_xml(heading_text))
            all_paragraphs.extend(paras)

            if progress_callback:
                progress_callback(i + 1, len(file_list), f.name)

        # Preserve section properties from the last file
        last_dir = tmp / "last"
        unpack_docx(file_list[-1], last_dir)
        sect_pr = get_sectPr_raw(last_dir / "word" / "document.xml")

        body_content = "\n".join(all_paragraphs)
        if sect_pr:
            body_content += f"\n{sect_pr}"

        new_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
mc:Ignorable="w14">
<w:body>
{body_content}
</w:body>
</w:document>'''

        with open(base_xml, 'w', encoding='utf-8') as f:
            f.write(new_xml)

        pack_docx(base_dir, output_path)
        return from_doc_count, from_filename_count


def format_chapter_headings(xml_path: Path) -> int:
    with open(xml_path, 'r', encoding='utf-8') as f:
        content = f.read()

    chapter_count = 0

    for pattern in CHAPTER_PATTERNS:
        core = pattern.strip('^$')
        regex = r'(<w:br/>)\s*(<w:t[^>]*>)(' + core + r')(</w:t>)\s*(<w:br/>)'

        def replace_heading(match):
            nonlocal chapter_count
            chapter_count += 1
            return f'''{match.group(1)}
</w:r>
</w:p>
<w:p>
<w:pPr>
<w:pStyle w:val="Heading1"/>
</w:pPr>
<w:r>
<w:t>{match.group(3)}</w:t>
</w:r>
</w:p>
<w:p>
<w:r>
{match.group(5)}'''

        content = re.sub(regex, replace_heading, content, flags=re.IGNORECASE)

    for pattern in CHAPTER_PATTERNS:
        core = pattern.strip('^$')
        regex2 = (
            r'(<w:t[^>]*>第[一二三四五六七八九十百千\d]+卷[^<]*</w:t>)\s*'
            r'(<w:br/>)\s*(<w:t[^>]*>)(' + core + r')(</w:t>)\s*(<w:br/>)'
        )

        def replace_volume(match):
            nonlocal chapter_count
            chapter_count += 1
            return f'''{match.group(1)}
</w:r>
</w:p>
<w:p>
<w:pPr>
<w:pStyle w:val="Heading1"/>
</w:pPr>
<w:r>
<w:t>{match.group(4)}</w:t>
</w:r>
</w:p>
<w:p>
<w:r>
{match.group(6)}'''

        content = re.sub(regex2, replace_volume, content, flags=re.IGNORECASE)

    with open(xml_path, 'w', encoding='utf-8') as f:
        f.write(content)

    return chapter_count


def process_combiner_files(sorted_files: list, output_path: Path, progress_callback=None):
    """
    Returns (from_doc_count, from_filename_count) — one H1 per file, guaranteed.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        merged_path = tmp / "merged.docx"
        from_doc, from_filename = merge_docx_files(sorted_files, merged_path, progress_callback)

        # Repack (no extra heading formatting needed — headings were built as H1 during merge)
        unpack_dir = tmp / "unpacked"
        unpack_docx(merged_path, unpack_dir)
        doc_xml = unpack_dir / "word" / "document.xml"

        if not doc_xml.exists():
            raise ValueError("No document.xml found — the .docx files may be corrupted.")

        pack_docx(unpack_dir, output_path)
        return from_doc, from_filename


def render_docx_combiner():
    st.title("📚 DocX Combiner & Formatter")
    st.markdown(
        "Upload multiple `.docx` files and this tool will automatically **sort them "
        "by episode/chapter number**, **merge them in the correct order**, and "
        "**format chapter headings** as Heading 1. "
        "If a file has no chapter heading inside it, a heading is created "
        "**from the filename** (e.g. `Show_42.docx` → `Chapter 42`)."
    )
    st.divider()

    uploaded_files = st.file_uploader(
        "Upload your .docx files",
        type=["docx"],
        accept_multiple_files=True,
        help="Select all the chapter/episode files at once. They will be sorted automatically.",
    )

    output_name = st.text_input(
        "Output file name",
        value="Combined",
        help="Your file will be saved as <name>_Combined_Formatted.docx",
    )
    st.divider()

    if uploaded_files:
        st.subheader(f"📂 {len(uploaded_files)} file(s) detected")
        file_preview = sorted(
            [(uf.name, extract_start_number(Path(uf.name))) for uf in uploaded_files],
            key=lambda x: x[1],
        )
        with st.expander("🔍 Preview detected merge order", expanded=True):
            for fname, num in file_preview:
                label = str(num) if num > 0 else "?"
                st.markdown(f"&nbsp;&nbsp;`[{label:>6}]`&nbsp;&nbsp; {fname}")
            st.markdown("")

        if st.button("🔗 Merge & Format", type="primary", use_container_width=True):
            progress_bar = st.progress(0, text="Preparing files...")
            result_area = st.empty()

            with tempfile.TemporaryDirectory() as tmp:
                tmp_path = Path(tmp)
                input_dir = tmp_path / "input"
                input_dir.mkdir()

                for uf in uploaded_files:
                    (input_dir / uf.name).write_bytes(uf.getvalue())

                docx_files = list(input_dir.glob("*.docx"))
                sorted_files = sort_batch_files(docx_files)

                safe_name = re.sub(r'[^\w\-_. ]', '_', output_name.strip() or "Combined")
                output_filename = f"{safe_name}_Combined_Formatted.docx"
                output_path = tmp_path / output_filename

                def update_progress(current, total, filename):
                    pct = int((current / total) * 80)
                    progress_bar.progress(pct, text=f"Merging ({current}/{total}): {filename}")

                try:
                    from_doc, from_filename = process_combiner_files(
                        sorted_files, output_path, update_progress
                    )
                    progress_bar.progress(100, text="Complete!")

                    with open(output_path, "rb") as f:
                        file_bytes = f.read()

                    result_area.success(
                        f"✅ Done! Merged **{len(sorted_files)} file(s)** · "
                        f"**{from_doc}** heading(s) taken from inside the doc · "
                        f"**{from_filename}** heading(s) built from filename. "
                        f"Exactly one H1 per chapter."
                    )
                    st.download_button(
                        label=f"⬇️ Download {output_filename}",
                        data=file_bytes,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary",
                    )

                except Exception as e:
                    progress_bar.empty()
                    result_area.error(f"❌ Something went wrong: {e}")
                    with st.expander("Show error details"):
                        st.code(traceback.format_exc())
    else:
        st.info("👆 Upload your `.docx` files above to get started.")

    st.divider()
    st.caption(
        "Supported chapter heading formats: &nbsp;"
        "`第X章` &nbsp;·&nbsp; `Chapter X` &nbsp;·&nbsp; `CHAPTER X` &nbsp;·&nbsp; "
        "`Ch. X` &nbsp;·&nbsp; `Episode X` &nbsp;·&nbsp; `卷X` &nbsp;·&nbsp; "
        "Chinese numeral chapters. &nbsp; If none are found inside a file, "
        "a heading is built from the filename (e.g. `Show_42.docx` → `Chapter 42`)."
    )


# ══════════════════════════════════════════════════════════════════════
#  TOOL 2 — BENCHMARK CONVERTER
# ══════════════════════════════════════════════════════════════════════

def extract_show_info_from_filename(filename: str):
    name = re.sub(r'\.docx$', '', filename, flags=re.IGNORECASE)
    name = re.sub(r'\s*\(\d+\)$', '', name)
    match = re.search(r'^(.+?)_(\d+)-(\d+)$', name)
    if match:
        return match.group(1), int(match.group(2)), int(match.group(3))
    alt = re.search(r'^(.+?)_(\d+)-(\d+)', name)
    if alt:
        return alt.group(1), int(alt.group(2)), int(alt.group(3))
    return name, 1, 500


def set_document_background_pagination(doc):
    try:
        settings_element = doc.settings.element
        bg_repag = settings_element.find(qn('w:displayBackgroundShape'))
        if bg_repag is None:
            bg_repag = OxmlElement('w:displayBackgroundShape')
            bg_repag.set(qn('w:val'), '1')
            settings_element.append(bg_repag)
        view_element = settings_element.find(qn('w:view'))
        if view_element is None:
            view_element = OxmlElement('w:view')
            view_element.set(qn('w:val'), 'print')
            settings_element.append(view_element)
    except Exception:
        pass


# ── Benchmark-Converter helpers ──────────────────────────────────────────────
# Strict episode/chapter regexes. A narrative sentence like "1. She walked in."
# will NOT match these because we require the Episode/Chapter/Ep/Ch keyword,
# or (in the style-based check below) an exact "Heading 1" style + bare number.
_EPISODE_REGEXES = [
    re.compile(r'^\s*(?:Episode|Chapter)\s*[-_:.]?\s*(\d+)\b', re.IGNORECASE),
    re.compile(r'^\s*(?:Ep|Ch)\.?\s*[-_:.]?\s*(\d+)\b',        re.IGNORECASE),
    re.compile(r'^\s*EP\s*[-_:.]?\s*(\d+)\b',                  re.IGNORECASE),
]

# Strips metadata like "WORD COUNT: 1512", "WORD COUNT - 1361", "Word count: 1350".
# Also catches the concatenated garbage case "THE WRONG ANAWORD COUNT - 1361".
_WORDCOUNT_STRIP = re.compile(r'\s*WORD\s*COUNT\s*[:\-–]?\s*\d+\s*$', re.IGNORECASE)


def _is_chapter_heading(text: str, style_name: str):
    """
    Decide if a paragraph is a chapter heading — returns (is_heading, ep_num).

    A paragraph counts as a chapter heading ONLY when either
      (a) its text matches a strict Episode/Chapter/Ep/Ch regex, OR
      (b) its style is EXACTLY 'Heading 1' AND its text looks like a
          chapter line (bare number, or number + separator + title).

    Heading 2 / Heading 3 / subtitle paragraphs in the source are NOT
    promoted — per Taggen's rule: "Only the chapter name or number
    should be formatted as a heading. No other content should be in a
    heading style."
    """
    if not text:
        return False, None

    for rx in _EPISODE_REGEXES:
        m = rx.match(text)
        if m:
            return True, int(m.group(1))

    if style_name == 'Heading 1':
        # "1 - Title" / "1: Title" / "1. Title"
        m = re.match(r'^\s*(\d+)\s*[-–—:.\s]', text)
        if m:
            return True, int(m.group(1))
        # bare number
        m = re.match(r'^\s*(\d+)\s*$', text)
        if m:
            return True, int(m.group(1))
        # Chinese markers
        if re.match(r'^\s*(?:第[\d一二三四五六七八九十百千]+[章卷])', text):
            return True, None

    return False, None


def _clean_heading_text(text: str, episode_num, fallback_num: int) -> str:
    """
    Normalise a chapter heading line for the benchmark output.

    - Strips trailing 'WORD COUNT …' junk.
    - Collapses internal whitespace.
    - Rewrites 'EPISODE 6 : THE PEONY RECORD' → 'Episode 6: The Peony Record'
      (single H1, no split).
    - Falls back to 'Episode N' if the source is just a bare number.
    """
    cleaned = _WORDCOUNT_STRIP.sub('', text).strip()
    cleaned = re.sub(r'\s{2,}', ' ', cleaned)

    num = episode_num if episode_num is not None else fallback_num

    m = re.match(
        r'^\s*(?:Episode|Chapter|Ep\.?|Ch\.?|EP)\s*[-_:.]?\s*\d+\s*'
        r'[-–—:.\s]*\s*(.*?)\s*$',
        cleaned,
        flags=re.IGNORECASE,
    )
    if m:
        subtitle = m.group(1).strip(' -–—:.')
        return f"Episode {num}: {subtitle}" if subtitle else f"Episode {num}"

    if cleaned and not cleaned[0].isdigit():
        return f"Episode {num}: {cleaned}"
    return f"Episode {num}"


def convert_single_file_to_benchmark(file_bytes: bytes, filename: str) -> tuple[bytes, str]:
    """
    Convert a .docx into the Taggen benchmark format.

    Per the Taggen guide:
      • Every chapter name/number must START with a Heading style.
      • ONLY the chapter name/number should be a heading — nothing else.
      • Filename: <show_name>_<first_ch>-<last_ch>.docx

    Therefore this writes EXACTLY ONE Heading 1 per detected chapter and
    demotes everything else (including source Heading 2 subtitles,
    'Word Count: 1671' lines, body text) to Normal paragraphs.
    """
    show_name, first_ep, last_ep = extract_show_info_from_filename(filename)

    doc     = Document(io.BytesIO(file_bytes))
    new_doc = Document()
    set_document_background_pagination(new_doc)

    for section in new_doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    styles = new_doc.styles
    try:
        styles['Heading 1']
    except KeyError:
        styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    current_episode   = first_ep
    paragraphs_to_add = []
    seen_any_heading  = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name         = para.style.name if para.style else ''
        is_heading, ep_num = _is_chapter_heading(text, style_name)

        if is_heading:
            heading_text    = _clean_heading_text(text, ep_num, current_episode)
            current_episode = (ep_num if ep_num is not None else current_episode) + 1
            paragraphs_to_add.append(('heading1', heading_text))
            seen_any_heading = True
        else:
            # EVERYTHING else → Normal (source Heading 2, Word Count lines, body).
            paragraphs_to_add.append(('normal', text))

    # Safety net: Taggen requires the file to start with a heading.
    if not seen_any_heading:
        paragraphs_to_add.insert(0, ('heading1', f"Episode {first_ep}"))

    for para_type, para_text in paragraphs_to_add:
        if para_type == 'heading1':
            p = new_doc.add_paragraph(para_text)
            p.style = new_doc.styles['Heading 1']
        else:
            new_doc.add_paragraph(para_text)

    output_filename = f"{show_name}_{first_ep}-{last_ep}.docx"
    buf = io.BytesIO()
    new_doc.save(buf)
    return buf.getvalue(), output_filename


def render_benchmark_converter():
    st.title("⚙️ Benchmark Converter")
    st.markdown(
        "Upload one or more `.docx` files — **one per show** — and this tool will "
        "convert each file into the **benchmark format** accepted by PocketFM's "
        "internal review tool."
    )
    st.markdown(
        "**Expected filename format:** `ShowName_FirstEp-LastEp.docx`  \n"
        "Example: `The_Man_They_Forgot_1-500.docx`"
    )
    st.divider()

    uploaded_files = st.file_uploader(
        "Upload .docx files (one per show)",
        type=["docx"],
        accept_multiple_files=True,
        key="bench_uploader",
    )
    st.divider()

    if not uploaded_files:
        st.info("👆 Upload your `.docx` files above to get started.")
        return

    st.subheader(f"📂 {len(uploaded_files)} file(s) ready to convert")
    with st.expander("🔍 Files detected", expanded=True):
        for uf in uploaded_files:
            show_name, first_ep, last_ep = extract_show_info_from_filename(uf.name)
            st.markdown(
                f"&nbsp;&nbsp;📄 **{uf.name}**  →  "
                f"Show: `{show_name}` &nbsp;|&nbsp; Episodes: `{first_ep}–{last_ep}`"
            )
        st.markdown("")

    if not st.button("🚀 Convert All Files", type="primary", use_container_width=True):
        return

    progress_bar = st.progress(0, text="Starting conversion...")
    errors  = []
    results = []

    for i, uf in enumerate(uploaded_files):
        progress_bar.progress(
            int(i / len(uploaded_files) * 100),
            text=f"Converting ({i + 1}/{len(uploaded_files)}): {uf.name}",
        )
        try:
            out_bytes, out_name = convert_single_file_to_benchmark(uf.getvalue(), uf.name)
            results.append((out_bytes, out_name))
        except Exception as e:
            errors.append((uf.name, str(e), traceback.format_exc()))

    progress_bar.progress(100, text="Done!")

    if errors:
        for fname, err_msg, tb in errors:
            st.error(f"❌ Failed to convert **{fname}**: {err_msg}")
            with st.expander(f"Error details — {fname}"):
                st.code(tb)

    if not results:
        return

    st.success(f"✅ Successfully converted **{len(results)}** file(s)!")
    st.markdown("---")

    if len(results) == 1:
        out_bytes, out_name = results[0]
        st.download_button(
            label=f"⬇️ Download {out_name}",
            data=out_bytes,
            file_name=out_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    else:
        st.markdown("**Download individual files:**")
        for out_bytes, out_name in results:
            st.download_button(
                label=f"⬇️ {out_name}",
                data=out_bytes,
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        st.markdown("**Or download everything at once:**")
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for out_bytes, out_name in results:
                zf.writestr(out_name, out_bytes)
        st.download_button(
            label=f"📦 Download all {len(results)} files as ZIP",
            data=zip_buf.getvalue(),
            file_name="benchmark_converted.zip",
            mime="application/zip",
            use_container_width=True,
            type="primary",
        )

    st.divider()
    st.caption(
        "Heading patterns detected: `Episode N` · `Chapter N` · `Ep. N` · `Ch. N` · `N.` · "
        "any paragraph already styled as Heading."
    )


# ══════════════════════════════════════════════════════════════════════
#  TOOL 3 — DOCUMENT TRANSLATOR  (Free Google Translate  OR  Gemini AI)
# ══════════════════════════════════════════════════════════════════════

GEMINI_BATCH_SIZE      = 25   # paragraphs per Gemini call (sweet spot for reliable parsing)
GEMINI_MAX_CONCURRENT  = 12   # max simultaneous Gemini API calls across ALL files
TRANS_FILE_WORKERS     = 5    # parallel files (both engines)
GTRANS_BATCH_SIZE      = 20   # paragraphs per free-API call
GTRANS_BATCH_WORKERS   = 4    # concurrent free-API batches per file (keep low to avoid throttling)


LANGUAGES = {
    "Chinese (Simplified)": "Chinese (Simplified)",
    "Chinese (Traditional)": "Chinese (Traditional)",
    "Korean": "Korean",
    "Japanese": "Japanese",
    "Arabic": "Arabic",
    "Spanish": "Spanish",
    "French": "French",
    "German": "German",
    "Hindi": "Hindi",
    "Portuguese": "Portuguese",
    "Russian": "Russian",
    "English": "English",
}

GTRANS_LANG_CODES = {
    "Chinese (Simplified)": "zh-CN",
    "Chinese (Traditional)": "zh-TW",
    "Korean": "ko",
    "Japanese": "ja",
    "Arabic": "ar",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Hindi": "hi",
    "Portuguese": "pt",
    "Russian": "ru",
    "English": "en",
}

GEMINI_MODELS = [
    "gemini-1.5-flash",
    "gemini-2.0-flash",
    "gemini-1.5-pro",
]


# ── Free Google Translate engine ─────────────────────────────────────────────

GTRANSLATE_URL = "https://translate.googleapis.com/translate_a/single"


def _gtrans_call(text: str, src_code: str, tgt_code: str) -> str:
    """Single call to the free Google Translate endpoint."""
    import requests as _req
    r = _req.get(
        GTRANSLATE_URL,
        params={"client": "gtx", "sl": src_code, "tl": tgt_code, "dt": "t", "q": text},
        timeout=20,
    )
    r.raise_for_status()
    data = r.json()
    return "".join(part[0] for part in data[0] if part[0])


def _gtrans_batch(texts: list, src_code: str, tgt_code: str) -> list:
    """
    Translate a batch via free Google Translate API.
    RAISES on failure — never silently returns original text.
    """
    if not texts:
        return texts
    try:
        result = _gtrans_call("\n".join(texts), src_code, tgt_code)
        parts  = [p.strip() for p in result.split("\n")]
        if len(parts) == len(texts):
            return parts
    except Exception as e:
        raise RuntimeError(
            f"Google Translate API call failed: {e}\n"
            "If this keeps happening, the free endpoint may be blocked from this server — "
            "try switching to Gemini mode instead."
        )

    # Fallback: try one-by-one
    results = []
    for t in texts:
        results.append(_gtrans_call(t, src_code, tgt_code) if t.strip() else t)
    return results


def _translate_docx_gtrans(file_bytes: bytes, src: str, tgt: str,
                            progress_cb=None) -> bytes:
    """Translate a .docx using the free Google Translate API (parallel batches)."""
    src_code = GTRANS_LANG_CODES.get(src, "auto")
    tgt_code = GTRANS_LANG_CODES.get(tgt, "en")

    src_doc   = Document(io.BytesIO(file_bytes))
    all_texts = [p.text for p in src_doc.paragraphs]
    non_empty = [(i, t) for i, t in enumerate(all_texts) if t.strip()]
    texts     = [t for _, t in non_empty]
    idx_map   = [i for i, _ in non_empty]

    translated_map = {}
    completed      = [0]
    lock           = threading.Lock()

    batches = [(s, texts[s:s + GTRANS_BATCH_SIZE]) for s in range(0, len(texts), GTRANS_BATCH_SIZE)]

    def run_batch(start, batch_texts):
        res = _gtrans_batch(batch_texts, src_code, tgt_code)
        with lock:
            for j, trans in enumerate(res):
                translated_map[idx_map[start + j]] = trans
            completed[0] += 1
            if progress_cb:
                progress_cb(completed[0], len(batches))

    with ThreadPoolExecutor(max_workers=GTRANS_BATCH_WORKERS) as ex:
        futs = [ex.submit(run_batch, s, b) for s, b in batches]
        for f in as_completed(futs):
            f.result()

    dst_doc = Document()
    for i, para in enumerate(src_doc.paragraphs):
        text     = translated_map.get(i, para.text)
        new_para = dst_doc.add_paragraph(text)
        try:
            if para.style and para.style.name:
                new_para.style = dst_doc.styles[para.style.name]
        except Exception:
            pass

    buf = io.BytesIO()
    dst_doc.save(buf)
    return buf.getvalue()


# ── Gemini engine ─────────────────────────────────────────────────────────────

def _gemini_translate_batch(texts: list, src: str, tgt: str, model) -> list:
    """
    Send a numbered list to Gemini, get back numbered translations.
    RAISES RuntimeError on failure — never silently returns original text.
    """
    if not texts:
        return texts

    numbered = "\n".join(f"{i+1}. {t}" for i, t in enumerate(texts))
    prompt = (
        f"You are a professional translator. Translate every numbered item below from {src} to {tgt}.\n"
        f"RULES:\n"
        f"- Output ONLY the translated lines, numbered identically (1. 2. 3. …)\n"
        f"- Never skip, merge, or add items — the count must stay exactly {len(texts)}\n"
        f"- Do not add any explanation, header, or extra text\n\n"
        f"{numbered}"
    )

    last_error = "Unknown error"
    for attempt in range(3):
        try:
            response = model.generate_content(prompt)
            raw      = response.text.strip()

            # Primary parse: "N. text" lines
            parsed = []
            for line in raw.splitlines():
                m = re.match(r'^\d+\.\s*(.+)', line.strip())
                if m:
                    parsed.append(m.group(1).strip())

            if len(parsed) == len(texts):
                return parsed

            # Fallback: non-empty, non-numeric-only lines
            lines = [l.strip() for l in raw.splitlines()
                     if l.strip() and not re.match(r'^\d+\.?\s*$', l.strip())]
            if len(lines) == len(texts):
                return lines

            last_error = (
                f"Gemini returned {len(parsed)} items, expected {len(texts)}. "
                f"Raw preview: {raw[:300]}"
            )
        except Exception as e:
            last_error = str(e)
            if attempt < 2:
                time.sleep(2)

    raise RuntimeError(f"Gemini translation failed after 3 attempts: {last_error}")


def _translate_docx_gemini(file_bytes: bytes, src: str, tgt: str,
                            model, progress_cb=None,
                            semaphore: threading.Semaphore = None) -> bytes:
    """
    Translate all paragraphs in a .docx using Gemini.
    Batches run in parallel, gated by a shared semaphore to respect rate limits.
    """
    src_doc   = Document(io.BytesIO(file_bytes))
    all_texts = [p.text for p in src_doc.paragraphs]

    non_empty      = [(i, t) for i, t in enumerate(all_texts) if t.strip()]
    texts          = [t for _, t in non_empty]
    idx_map        = [i for i, _ in non_empty]
    translated_map = {}
    completed      = [0]
    map_lock       = threading.Lock()

    batches = [(s, texts[s:s + GEMINI_BATCH_SIZE])
               for s in range(0, len(texts), GEMINI_BATCH_SIZE)]
    total   = len(batches)

    def run_batch(start, batch):
        ctx = semaphore if semaphore else threading.Semaphore(GEMINI_MAX_CONCURRENT)
        with ctx:
            result = _gemini_translate_batch(batch, src, tgt, model)
        with map_lock:
            for j, trans_text in enumerate(result):
                translated_map[idx_map[start + j]] = trans_text
            completed[0] += 1
            if progress_cb:
                progress_cb(completed[0], total)

    with ThreadPoolExecutor(max_workers=GEMINI_MAX_CONCURRENT) as ex:
        futs = [ex.submit(run_batch, s, b) for s, b in batches]
        for f in as_completed(futs):
            f.result()

    dst_doc = Document()
    for i, para in enumerate(src_doc.paragraphs):
        text     = translated_map.get(i, para.text)
        new_para = dst_doc.add_paragraph(text)
        try:
            if para.style and para.style.name:
                new_para.style = dst_doc.styles[para.style.name]
        except Exception:
            pass

    buf = io.BytesIO()
    dst_doc.save(buf)
    return buf.getvalue()


def _fetch_files_from_drive(folder_url: str) -> dict:
    """
    Download all .docx files from a public Google Drive folder.
    Returns {filename: bytes}.
    """
    if not GDOWN_AVAILABLE:
        raise RuntimeError("gdown is not installed. Cannot fetch from Drive.")

    tmp_dir = tempfile.mkdtemp()
    try:
        gdown.download_folder(
            url=folder_url,
            output=tmp_dir,
            quiet=True,
            use_cookies=False,
        )
        result = {}
        for root, _, filenames in os.walk(tmp_dir):
            for fn in filenames:
                if fn.lower().endswith(".docx"):
                    with open(os.path.join(root, fn), "rb") as f:
                        result[fn] = f.read()
        return result
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def render_document_translator():
    st.title("🌐 Document Translator")
    st.markdown(
        "Translate `.docx` files — upload from your computer or paste a Google Drive folder link. "
        "All files process in parallel and download as a single ZIP."
    )
    st.divider()

    for key, default in [("tr_results", {}), ("tr_errors", {}), ("tr_api_key", "")]:
        if key not in st.session_state:
            st.session_state[key] = default

    engine = st.radio(
        "Translation engine",
        options=["🆓  Free — Google Translate (no key needed)",
                 "✨  Gemini AI — higher quality (API key required)"],
        key="tr_engine",
        horizontal=True,
    )
    use_gemini = engine.startswith("✨")

    model_name = GEMINI_MODELS[0]

    if use_gemini:
        if not GEMINI_AVAILABLE:
            st.error("⛔ `google-generativeai` package missing. Add it to requirements.txt and redeploy.")
            return

        st.markdown(
            """
            <div style="
                background: linear-gradient(135deg, #667eea22, #764ba222);
                border: 2px solid #667eea;
                border-radius: 12px;
                padding: 18px 20px 10px 20px;
                margin: 12px 0 4px 0;
            ">
            <p style="margin:0 0 6px 0; font-weight:700; font-size:1rem;">
                🔑 &nbsp;Gemini API Key <span style="color:#667eea">(required)</span>
            </p>
            """,
            unsafe_allow_html=True,
        )
        api_key_input = st.text_input(
            label="gemini_key_label",
            label_visibility="collapsed",
            type="password",
            placeholder="AIzaSy…  — paste your key here",
            value=st.session_state.tr_api_key,
            key="tr_api_key_input",
        )
        st.markdown(
            """
            <p style="margin:4px 0 0 0; font-size:0.8rem; color:#888;">
            Get a free key at
            <a href="https://aistudio.google.com/apikey" target="_blank">aistudio.google.com/apikey</a>
            &nbsp;·&nbsp; gemini-1.5-flash is free up to 1,500 requests/day
            </p></div>
            """,
            unsafe_allow_html=True,
        )
        if api_key_input:
            st.session_state.tr_api_key = api_key_input

        st.markdown("<div style='margin-top:10px'></div>", unsafe_allow_html=True)
        model_name = st.selectbox("Gemini model", GEMINI_MODELS, index=0, key="tr_model")
    else:
        st.caption("Uses Google Translate's free public endpoint — no account or key needed.")

    st.divider()

    col1, col2, col3 = st.columns([2, 0.6, 2])
    with col1:
        src_name = st.selectbox("Source language", list(LANGUAGES.keys()),
                                index=0, key="tr_src")
    with col2:
        st.markdown("<br><div style='text-align:center;font-size:22px'>→</div>",
                    unsafe_allow_html=True)
    with col3:
        tgt_name = st.selectbox("Target language", list(LANGUAGES.keys()),
                                index=11, key="tr_tgt")

    src_lang = LANGUAGES[src_name]
    tgt_lang = LANGUAGES[tgt_name]

    st.divider()

    tab_upload, tab_drive = st.tabs(["📁  Upload Files from Computer",
                                     "🔗  Google Drive Folder Link"])

    with tab_upload:
        uploaded_files = st.file_uploader(
            "Select one or more .docx files",
            type=["docx"],
            accept_multiple_files=True,
            key="tr_uploader",
            help="Hold Cmd/Ctrl to select multiple files at once",
        )

    with tab_drive:
        st.markdown("")
        drive_link = st.text_input(
            "Paste Google Drive folder link",
            placeholder="https://drive.google.com/drive/folders/…",
            key="tr_drive_link",
        )
        st.caption(
            "⚠️ Folder must be shared as **'Anyone with the link can view'** — "
            "private folders require the Colab notebook instead."
        )

    st.divider()

    has_input = bool(uploaded_files) or bool(drive_link)
    has_key   = bool(st.session_state.tr_api_key) if use_gemini else True
    ready     = has_input and has_key

    if use_gemini and not has_key:
        st.warning("⬆️ Enter your Gemini API key above to enable translation.")
    elif not has_input:
        st.info("⬆️ Upload files or paste a Drive folder link to get started.")

    if use_gemini and has_key:
        if st.button("🔬 Test API Key", key="tr_test"):
            with st.spinner("Testing Gemini API key…"):
                try:
                    genai.configure(api_key=st.session_state.tr_api_key)
                    test_model    = genai.GenerativeModel(model_name)
                    test_response = test_model.generate_content(
                        f"Translate this from Chinese (Simplified) to English. "
                        f"Reply with ONLY the translation, nothing else: 你好，世界"
                    )
                    result_text = test_response.text.strip()
                    st.success(f"✅ API key works! Test translation: **{result_text}**")
                except Exception as e:
                    st.error(f"❌ API key test failed: {e}")

    if st.button("🚀  Translate All & Download ZIP", type="primary",
                 use_container_width=True, disabled=not ready, key="tr_go"):

        st.session_state.tr_results = {}
        st.session_state.tr_errors  = {}

        gemini_model     = None
        gemini_semaphore = None
        if use_gemini:
            genai.configure(api_key=st.session_state.tr_api_key)
            gemini_model     = genai.GenerativeModel(model_name)
            gemini_semaphore = threading.Semaphore(GEMINI_MAX_CONCURRENT)

        file_bytes = {}
        if uploaded_files:
            for f in uploaded_files:
                file_bytes[f.name] = f.read()

        if drive_link:
            with st.spinner("📥 Fetching files from Google Drive…"):
                try:
                    drive_files = _fetch_files_from_drive(drive_link)
                    if not drive_files:
                        st.error("No .docx files found in that Drive folder.")
                    else:
                        file_bytes.update(drive_files)
                        st.success(f"✅ Downloaded {len(drive_files)} file(s) from Drive.")
                except Exception as e:
                    st.error(f"❌ Could not fetch from Drive: {e}")

        if not file_bytes:
            st.stop()

        engine_label = f"Gemini · {model_name}" if use_gemini else "Google Translate (free)"
        st.info(f"Translating **{len(file_bytes)} file(s)** — "
                f"{src_lang} → {tgt_lang} · {engine_label}")

        progress_state = {
            name: {"done": 0, "total": 1, "status": "queued"}
            for name in file_bytes
        }
        results    = {}
        errors     = {}
        outer_lock = threading.Lock()

        ui = {}
        for name in file_bytes:
            st.markdown(f"**{name}**")
            ui[name] = {"bar": st.progress(0), "status": st.empty()}
            ui[name]["status"].text("⌛ Queued…")
        st.divider()

        def file_worker(name, content):
            progress_state[name]["status"] = "running"

            def on_batch(done, total):
                progress_state[name].update(done=done, total=total)

            try:
                if use_gemini:
                    translated = _translate_docx_gemini(
                        content, src_lang, tgt_lang, gemini_model,
                        progress_cb=on_batch, semaphore=gemini_semaphore
                    )
                else:
                    translated = _translate_docx_gtrans(
                        content, src_lang, tgt_lang, progress_cb=on_batch
                    )
                with outer_lock:
                    results[name] = translated
                progress_state[name]["status"] = "done"
            except Exception as e:
                with outer_lock:
                    errors[name] = str(e)
                progress_state[name]["status"] = "error"

        with ThreadPoolExecutor(max_workers=TRANS_FILE_WORKERS) as executor:
            futures = {
                executor.submit(file_worker, n, c): n
                for n, c in file_bytes.items()
            }
            while not all(f.done() for f in futures):
                for name, state in progress_state.items():
                    pct = state["done"] / max(state["total"], 1)
                    s   = state["status"]
                    if s == "queued":
                        ui[name]["status"].text("⌛ Queued…")
                    elif s == "running":
                        ui[name]["bar"].progress(pct)
                        if pct >= 1.0:
                            ui[name]["status"].text("📝 Finalizing document…")
                        else:
                            ui[name]["status"].text(
                                f"⏳ {state['done']}/{state['total']} batches ({int(pct*100)}%)"
                            )
                    elif s == "done":
                        ui[name]["bar"].progress(1.0)
                        ui[name]["status"].text("✅ Done!")
                    elif s == "error":
                        ui[name]["status"].text(f"❌ {errors.get(name, 'unknown error')}")
                time.sleep(0.3)

        for name, state in progress_state.items():
            if state["status"] == "done":
                ui[name]["bar"].progress(1.0)
                ui[name]["status"].text("✅ Done!")
            elif state["status"] == "error":
                ui[name]["status"].text(f"❌ {errors.get(name, 'unknown error')}")

        st.session_state.tr_results = results
        st.session_state.tr_errors  = errors

        st.rerun()

    if st.session_state.get("tr_results"):
        results = st.session_state.tr_results
        errors  = st.session_state.tr_errors

        st.subheader("⬇️ Download Translated Files")

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in results.items():
                zf.writestr(name.replace(".docx", "_EN.docx"), data)
        zip_buf.seek(0)

        st.download_button(
            label=f"📦  Download ALL {len(results)} file(s) as ZIP",
            data=zip_buf.getvalue(),
            file_name="translated_docs.zip",
            mime="application/zip",
            type="primary",
            use_container_width=True,
            key="tr_zip",
        )

        if len(results) > 1:
            st.caption("Or download individual files:")
            cols = st.columns(min(len(results), 3))
            for i, (name, data) in enumerate(results.items()):
                out_name = name.replace(".docx", "_EN.docx")
                with cols[i % len(cols)]:
                    st.download_button(
                        label=f"⬇️ {out_name}",
                        data=data,
                        file_name=out_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"tr_dl_{name}",
                    )

        if errors:
            st.warning(f"⚠️ {len(errors)} file(s) failed: {', '.join(errors.keys())}")

    st.divider()
    st.caption(
        "🆓 Free mode: Google Translate public endpoint, 20 paragraphs/batch, 10 parallel batches · "
        "✨ Gemini mode: numbered-list batches of 30 paragraphs, higher accuracy"
    )


# ══════════════════════════════════════════════════════════════════════
#  ROUTER
# ══════════════════════════════════════════════════════════════════════

if selected_tool == "📚 DocX Combiner":
    render_docx_combiner()
elif selected_tool == "⚙️ Benchmark Converter":
    render_benchmark_converter()
else:
    render_document_translator()
